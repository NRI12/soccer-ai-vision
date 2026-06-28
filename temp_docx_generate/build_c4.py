"""Build chuong4_thuc_nghiem.docx from markdown via pandoc + python-docx post-processing."""
import subprocess, sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_LINE_SPACING

ROOT = Path(__file__).parent
import sys as _sys
_target = _sys.argv[1] if len(_sys.argv) > 1 else "c5"

if _target == "mo_dau":
    MD  = ROOT / "mo_dau.md"
    OUT = ROOT / "mo_dau.docx"
elif _target == "c1":
    MD  = ROOT / "chuong1_tong_quan.md"
    OUT = ROOT / "chuong1_tong_quan.docx"
else:
    MD  = ROOT / "chuong5_ket_luan.md"
    OUT = ROOT / "chuong5_ket_luan.docx"

REF  = ROOT / "reference_c4.docx"

# ── 1. pandoc ────────────────────────────────────────────────────────────────
cmd = [
    "pandoc", str(MD),
    "-o", str(OUT),
    f"--reference-doc={REF}",
    "--mathml",
]
print("Running pandoc…")
r = subprocess.run(cmd, capture_output=True, text=True, cwd=str(ROOT))
if r.returncode != 0:
    print("STDERR:", r.stderr)
    sys.exit(r.returncode)
print("pandoc OK ->", OUT)

# ── 2. python-docx post-processing ───────────────────────────────────────────
doc = Document(str(OUT))

def has_equation(para):
    return (para._p.find(qn("m:oMathPara")) is not None or
            para._p.find(qn("m:oMath"))     is not None)

# Page margins (L=3cm, R=2cm, T=2.5cm, B=3cm)
for section in doc.sections:
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.0)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(3.0)

# Line spacing + font normalisation
for para in doc.paragraphs:
    pf = para.paragraph_format

    if has_equation(para):
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing      = 1.0          # float, NOT Pt()
        pf.space_before      = Pt(8)
        pf.space_after       = Pt(8)
        para.alignment       = WD_ALIGN_PARAGRAPH.CENTER
    else:
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing      = 1.3          # float, NOT Pt()

    # Ensure Times New Roman 13pt for all runs
    for run in para.runs:
        if run.font.name is None or run.font.name == "":
            run.font.name = "Times New Roman"
        if run.font.size is None:
            run.font.size = Pt(13)

# Table font normalisation
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                pf = para.paragraph_format
                pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                pf.line_spacing      = 1.15
                for run in para.runs:
                    if not run.font.name:
                        run.font.name = "Times New Roman"
                    if run.font.size is None:
                        run.font.size = Pt(12)

doc.save(str(OUT))
print("Post-processing done ->", OUT)
size_kb = OUT.stat().st_size / 1024
print(f"File size: {size_kb:.1f} KB")
